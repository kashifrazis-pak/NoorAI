from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Colour palette ────────────────────────────────────────────────────────────
GREEN_DARK   = RGBColor(0x1A, 0x5C, 0x38)   # deep Islamic green
GREEN_MID    = RGBColor(0x2E, 0x86, 0x48)   # medium green
GREEN_LIGHT  = RGBColor(0xD6, 0xEE, 0xDB)   # very light green (table header bg)
GOLD         = RGBColor(0xC9, 0xA0, 0x2A)
GREY_TEXT    = RGBColor(0x44, 0x44, 0x44)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4 if level == 1 else 2)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=GREEN_DARK)
        # underline via border bottom
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1A5C38')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=GREEN_MID)
    elif level == 3:
        set_run_font(run, size=11, bold=True, color=GOLD)
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_run_font(run, size=11, color=GREY_TEXT)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=GREY_TEXT)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=GREY_TEXT)
    return p

def add_note(text, label="Note"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    lbl = p.add_run(f"  {label}: ")
    set_run_font(lbl, size=10, bold=True, color=GOLD)
    body = p.add_run(text)
    set_run_font(body, size=10, italic=True, color=GREY_TEXT)
    # light left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '8')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'C9A02A')
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        set_run_font(run, size=10, bold=header, color=WHITE if header else GREY_TEXT)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        if header:
            shade_cell(cell, '1A5C38')
        elif i == 0:
            shade_cell(cell, 'E8F5EC')

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("NoorAI")
set_run_font(r, size=36, bold=True, color=GREEN_DARK)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("AI-Powered Islamic Q&A Application")
set_run_font(r2, size=16, color=GREEN_MID)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tagline.add_run('"Answers rooted in the Quran and authentic Sunnah"')
set_run_font(r3, size=12, italic=True, color=GOLD)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = meta.add_run("Technical Build Plan  |  Version 1.0  |  May 2026")
set_run_font(r4, size=10, color=RGBColor(0x88, 0x88, 0x88))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("1. Executive Summary")
add_body(
    "NoorAI is a multilingual, cross-platform Islamic Q&A application that uses a large language "
    "model (LLM) to answer questions posed by Muslims around the world. Every response is grounded "
    "exclusively in two primary sources: the Holy Quran and the authentic (sahih/hasan) Hadith "
    "literature. The application will be available as a progressive web app (PWA) on desktop browsers "
    "and as native-quality mobile apps on iOS and Android, ensuring accessibility for the global "
    "Muslim community regardless of device."
)
add_body(
    "The system will cite specific Quran verses (Surah:Ayah) and Hadith references (book, volume, "
    "hadith number, narrator chain grade) alongside every answer, allowing users to verify responses "
    "directly in authoritative sources. A human scholar review layer will be built in from day one "
    "to catch edge cases and maintain trust."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. APP NAME & IDENTITY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("2. App Name & Identity")
add_body("Suggested working name: NoorAI  (Noor = Light in Arabic, symbolising guidance from divine sources)")
add_body("The final name, logo, and branding should be reviewed by Islamic scholars and community representatives before launch.")

# ═══════════════════════════════════════════════════════════════════════════════
# 3. GOALS & NON-GOALS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("3. Goals & Non-Goals")
add_heading("3.1  Goals", level=2)
add_bullet("Provide accurate, source-cited answers to Islamic questions in multiple languages.")
add_bullet("Draw responses exclusively from the Quran and the authentic Hadith collections (Kutub al-Sittah, Muwatta Imam Malik, Musnad Ahmad).")
add_bullet("Display Arabic text of cited Quran verses and Hadith alongside translations.")
add_bullet("Support multilingual input and output (English, Arabic, Urdu, and extensible to others).")
add_bullet("Work seamlessly on web browsers, iOS, and Android.")
add_bullet("Allow users to rate answers and flag potential errors for scholar review.")
add_bullet("Be accessible to scholars who can verify and annotate AI responses.")

add_heading("3.2  Non-Goals (what the app will NOT do)", level=2)
add_bullet("Issue formal fatwas or religious rulings — responses are educational, not authoritative legal verdicts.")
add_bullet("Reference weak (daif) or fabricated (mawdu) hadith in answers.")
add_bullet("Rely on non-primary sources such as books of fiqh, tafsir books, or scholarly opinions as primary data sources (these may be referenced as supplementary context only).")
add_bullet("Replace human Islamic scholars — the app positions itself as a learning aid, not an authority.")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. AUTHENTIC HADITH SOURCES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("4. Authentic Hadith Data Sources")
add_body(
    "The following authoritative collections will be ingested into the knowledge base. Only hadith "
    "graded Sahih (authentic) or Hasan (good) by classical Hadith scholars will be used in AI responses. "
    "Daif (weak) narrations will be stored but flagged and excluded from answer generation."
)

# Table
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.columns[0].width = Inches(2.4)
tbl.columns[1].width = Inches(1.6)
tbl.columns[2].width = Inches(3.2)
add_table_row(tbl, ["Collection", "Approx. Hadith", "Notes"], header=True)
rows_data = [
    ("Sahih al-Bukhari",       "7,563",  "Highest authenticity — most widely trusted collection"),
    ("Sahih Muslim",           "7,500",  "Second most authentic; companion to Bukhari"),
    ("Sunan Abu Dawud",        "5,274",  "Strong focus on legal/fiqh matters"),
    ("Jami al-Tirmidhi",       "3,956",  "Grades each hadith; includes daif — filter required"),
    ("Sunan al-Nasa'i",        "5,761",  "Strict in grading; less daif than others"),
    ("Sunan Ibn Majah",        "4,341",  "Part of the six; requires grading filter"),
    ("Muwatta Imam Malik",     "1,720",  "Earliest major collection; highly authenticated"),
    ("Musnad Ahmad ibn Hanbal","27,000+","Largest collection; grading filter essential"),
    ("Quran (full text)",      "6,236 ayat", "Arabic + translations in EN, UR, FR, TR and more"),
]
for r in rows_data:
    add_table_row(tbl, list(r))

doc.add_paragraph()
add_note(
    "All Hadith data is available via open-source APIs such as sunnah.com API and HadithAPI.com, "
    "or can be sourced from structured datasets on GitHub (e.g. hadith-json). Quran data is "
    "available via quran.com API (Al-Quran Cloud)."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CORE FEATURES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("5. Core Features")

add_heading("5.1  Question & Answer Engine", level=2)
add_bullet("Free-text question input in any supported language.")
add_bullet("AI response always cites specific Quran verses (e.g., Surah Al-Baqarah 2:286) and/or Hadith references (e.g., Sahih Bukhari, Book 2, Hadith 47).")
add_bullet("Arabic original text displayed alongside the translation for every citation.")
add_bullet("Confidence indicator shown on each response (High / Medium — escalate to scholar review).")
add_bullet("'I don't know' fallback when no relevant, authentic source can be found — the app must never fabricate references.")

add_heading("5.2  Multilingual Support", level=2)
add_bullet("Phase 1 languages: English, Arabic, Urdu.")
add_bullet("Phase 2 languages: French, Turkish, Indonesian, Bangla, Malay.")
add_bullet("Language auto-detection on input; user can override from settings.")
add_bullet("All Quran citations displayed in Arabic with translation in the user's selected language.")

add_heading("5.3  Source Viewer", level=2)
add_bullet("Tap any citation to open a full-text source card showing the complete verse or hadith.")
add_bullet("Deep-link to sunnah.com or quran.com for further reading.")
add_bullet("Option to copy citation text for sharing.")

add_heading("5.4  Topic Categories", level=2)
add_bullet("Prayer (Salah), Fasting (Sawm), Charity (Zakat), Pilgrimage (Hajj).")
add_bullet("Family & Marriage, Business & Finance (Halal/Haram), Food & Drink.")
add_bullet("Faith & Belief (Aqeedah), Morals & Character, Death & Afterlife.")
add_bullet("Quran Recitation & Memorisation guidance.")
add_bullet("Daily Duas (supplications) and Sunnah acts.")

add_heading("5.5  Scholar Review Layer", level=2)
add_bullet("Flagging system: users can flag any answer as potentially incorrect.")
add_bullet("Scholar dashboard: verified Islamic scholars can review flagged answers, add annotations, or mark answers as approved.")
add_bullet("Approved answers are cached and served directly on repeat questions.")

add_heading("5.6  User Accounts & History", level=2)
add_bullet("Optional account creation (email or Google/Apple sign-in).")
add_bullet("Save favourite Q&A pairs for offline access.")
add_bullet("Question history with ability to re-run or share.")
add_bullet("Guest mode available — no account required for basic use.")

add_heading("5.7  Daily Reminder & Hadith of the Day", level=2)
add_bullet("Optional push notification with a daily authentic hadith.")
add_bullet("Daily Quran verse with translation.")
add_bullet("Ramadan-specific features: prayer times, suhoor/iftar countdowns, tarawih recitation guide.")

# ═══════════════════════════════════════════════════════════════════════════════
# 6. TECH STACK
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("6. Recommended Technology Stack")

add_heading("6.1  Frontend (Web + Mobile)", level=2)
add_body(
    "Use React Native with Expo for mobile (iOS + Android from a single codebase), "
    "and Next.js for the web progressive web app (PWA). Share business logic and API calls "
    "between both via a shared TypeScript library."
)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
tbl2.columns[0].width = Inches(1.8)
tbl2.columns[1].width = Inches(2.2)
tbl2.columns[2].width = Inches(3.2)
add_table_row(tbl2, ["Layer", "Technology", "Reason"], header=True)
fe_rows = [
    ("Mobile App",   "React Native + Expo",   "Single codebase for iOS & Android; native performance"),
    ("Web PWA",      "Next.js (React)",        "SEO, SSR, installable PWA; shared components with mobile"),
    ("Styling",      "NativeWind / Tailwind",  "Consistent design tokens across platforms"),
    ("State Mgmt",   "Zustand or Redux Toolkit","Lightweight, predictable state management"),
    ("Internationalisation", "i18next + react-i18next", "RTL support for Arabic/Urdu; dynamic locale loading"),
    ("Arabic/RTL",   "Built-in RN RTL + CSS",  "Proper right-to-left layout for Arabic content"),
    ("Offline",      "React Query + SQLite",   "Cache answers for offline access; background sync"),
]
for r in fe_rows:
    add_table_row(tbl2, list(r))

doc.add_paragraph()
add_heading("6.2  Backend & API", level=2)

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
tbl3.columns[0].width = Inches(1.8)
tbl3.columns[1].width = Inches(2.2)
tbl3.columns[2].width = Inches(3.2)
add_table_row(tbl3, ["Layer", "Technology", "Reason"], header=True)
be_rows = [
    ("API Server",      "Node.js + Fastify or Python + FastAPI", "Fast async APIs; Python preferred for AI/ML pipeline integration"),
    ("Database",        "PostgreSQL",                            "Relational store for users, Q&A cache, feedback, scholar annotations"),
    ("Vector Database", "Pinecone or pgvector (PostgreSQL ext)", "Semantic search over Quran & Hadith embeddings"),
    ("Cache",           "Redis",                                 "Cache frequent Q&A pairs; session tokens"),
    ("Auth",            "Supabase Auth or Auth0",                "Email, Google, Apple sign-in; JWT tokens"),
    ("File Storage",    "AWS S3 or Supabase Storage",            "Audio recitations, profile images"),
    ("Search",          "Meilisearch or Elasticsearch",          "Full-text keyword search across Quran & Hadith"),
]
for r in be_rows:
    add_table_row(tbl3, list(r))

doc.add_paragraph()
add_heading("6.3  AI / LLM Layer", level=2)

tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
tbl4.columns[0].width = Inches(1.8)
tbl4.columns[1].width = Inches(2.2)
tbl4.columns[2].width = Inches(3.2)
add_table_row(tbl4, ["Component", "Technology", "Notes"], header=True)
ai_rows = [
    ("Primary LLM",          "Claude (Anthropic API)",    "Excellent instruction-following; citation adherence; multilingual"),
    ("Embeddings",           "OpenAI text-embedding-3 or Cohere multilingual", "Generate semantic vectors for Quran & Hadith corpus"),
    ("RAG Orchestration",    "LangChain or LlamaIndex",   "Retrieval-Augmented Generation pipeline linking LLM to knowledge base"),
    ("Prompt Engineering",   "Custom system prompts",     "Strict grounding instructions — see Section 8"),
    ("Language Detection",   "langdetect + LLM fallback", "Auto-detect user's language; route to correct translation index"),
    ("Hallucination Guard",  "Citation verification layer","Post-process every response to confirm every cited reference exists in DB"),
]
for r in ai_rows:
    add_table_row(tbl4, list(r))

doc.add_paragraph()
add_heading("6.4  Infrastructure & DevOps", level=2)

tbl5 = doc.add_table(rows=1, cols=2)
tbl5.style = 'Table Grid'
tbl5.columns[0].width = Inches(2.5)
tbl5.columns[1].width = Inches(4.75)
add_table_row(tbl5, ["Component", "Choice"], header=True)
infra_rows = [
    ("Cloud Provider",       "AWS or GCP (choose one; both support all required services)"),
    ("Container Orchestration", "Docker + AWS ECS (Fargate) or Cloud Run for serverless scaling"),
    ("CI/CD",                "GitHub Actions — automated test, lint, deploy on merge to main"),
    ("Monitoring",           "Datadog or Sentry for error tracking; LLM response logging dashboard"),
    ("App Store Deployment", "Expo EAS Build — automates iOS TestFlight & Android Play Store builds"),
    ("Domain & CDN",         "Cloudflare — CDN, DDoS protection, edge caching"),
]
for r in infra_rows:
    add_table_row(tbl5, list(r))

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ARCHITECTURE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("7. System Architecture Overview")
add_body(
    "The application follows a Retrieval-Augmented Generation (RAG) architecture. "
    "When a user submits a question, the system does NOT rely on the LLM's training data alone. "
    "Instead it retrieves the most relevant Quran verses and Hadith from a curated, pre-indexed "
    "vector database, then passes that retrieved context to the LLM together with a strict system "
    "prompt instructing it to answer only from the provided context."
)

add_heading("7.1  Request Flow (step by step)", level=2)
add_numbered("User types a question in the mobile/web app (any supported language).")
add_numbered("Frontend sends the question to the backend API via HTTPS.")
add_numbered("Language Detection Service identifies the input language.")
add_numbered("The question is converted to a semantic embedding vector using the Embeddings Model.")
add_numbered("The vector is used to query the Vector Database (Pinecone/pgvector), returning the top-K most semantically similar Quran verses and Hadith.")
add_numbered("Retrieved passages (with Arabic text, translation, full citation reference) are assembled into a context block.")
add_numbered("The Context Block + User Question are passed to the LLM (Claude) with a strict grounding system prompt.")
add_numbered("The LLM generates an answer citing only the provided sources.")
add_numbered("A Citation Verification Layer checks that every reference in the response exists in the database — if not, the response is rejected and regenerated.")
add_numbered("The verified response is stored in the PostgreSQL cache and returned to the frontend.")
add_numbered("Frontend renders the answer with formatted Arabic citations and tappable source cards.")

add_heading("7.2  Knowledge Base Indexing Pipeline", level=2)
add_bullet("Ingest: Parse all Quran and Hadith datasets from open-source APIs/JSON files into a structured PostgreSQL schema.")
add_bullet("Embed: Run each verse/hadith through the embedding model to produce a vector representation. Store the vector alongside the metadata (collection, book, chapter, hadith number, grade, Arabic text, translations).")
add_bullet("Index: Load all vectors into Pinecone or pgvector for fast approximate nearest-neighbour (ANN) search.")
add_bullet("Filter: At query time, exclude any hadith not graded Sahih or Hasan. Apply Quran-only filter when questions are explicitly about Quranic guidance.")
add_bullet("Refresh: Re-embed if translation updates or new scholarly gradings are published.")

# ═══════════════════════════════════════════════════════════════════════════════
# 8. AI SYSTEM PROMPT STRATEGY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("8. AI System Prompt Strategy (Critical)")
add_body(
    "The quality and trustworthiness of NoorAI depends almost entirely on the system prompt "
    "given to the LLM. The following instructions must be hardcoded and never overrideable by "
    "user input:"
)
add_bullet("You are an Islamic knowledge assistant. Your role is to help Muslims understand their faith.")
add_bullet("You MUST answer ONLY from the Quran verses and Hadith passages provided in the context block below. Do NOT use any knowledge from your training data that is not reflected in the provided context.")
add_bullet("Every answer MUST cite the exact source: for Quran use [Surah Name, Chapter:Verse], for Hadith use [Collection Name, Book X, Hadith Y].")
add_bullet("If the provided context does not contain enough information to answer the question confidently, respond: 'I was unable to find a direct answer in the available Quran and authentic Hadith sources. Please consult a qualified Islamic scholar.'")
add_bullet("Never issue a fatwa or definitive religious ruling. Frame answers as: 'According to [source]...' or 'The Prophet (PBUH) said...'")
add_bullet("Always include the original Arabic text of any cited verse or hadith, followed by its translation.")
add_bullet("Do not speculate, extrapolate, or fill gaps with general knowledge. Stick strictly to the retrieved context.")
add_bullet("Respond in the same language as the user's question unless they ask otherwise.")

add_note(
    "This RAG + strict grounding approach is the primary safeguard against hallucinated references. "
    "The citation verification post-processing step (Section 7.1, step 9) provides a second layer of protection."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. DATABASE SCHEMA (HIGH LEVEL)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("9. Database Schema (High-Level)")

add_heading("9.1  Core Tables", level=2)
tables_desc = [
    ("quran_verses",     "id, surah_number, surah_name_ar, surah_name_en, ayah_number, arabic_text, translations (JSONB), embedding_vector"),
    ("hadith",           "id, collection, book_number, hadith_number, grade (sahih/hasan/daif), arabic_text, translations (JSONB), narrator_chain, embedding_vector"),
    ("users",            "id, email, display_name, preferred_language, created_at, role (user/scholar/admin)"),
    ("questions",        "id, user_id (nullable), question_text, language, created_at"),
    ("answers",          "id, question_id, answer_text, citations (JSONB), confidence, model_version, verified_by_scholar, created_at"),
    ("flags",            "id, answer_id, user_id, reason, status (pending/reviewed), scholar_notes"),
    ("saved_items",      "id, user_id, answer_id, created_at"),
    ("daily_content",    "id, type (hadith/verse), content_id, scheduled_date, language"),
]
tbl6 = doc.add_table(rows=1, cols=2)
tbl6.style = 'Table Grid'
tbl6.columns[0].width = Inches(1.8)
tbl6.columns[1].width = Inches(5.45)
add_table_row(tbl6, ["Table", "Key Fields"], header=True)
for r in tables_desc:
    add_table_row(tbl6, list(r))

# ═══════════════════════════════════════════════════════════════════════════════
# 10. API ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("10. API Endpoints (REST)")

endpoints = [
    ("POST", "/api/v1/ask",                 "Submit a question; returns AI-generated answer with citations"),
    ("GET",  "/api/v1/answers/:id",         "Retrieve a specific answer by ID"),
    ("POST", "/api/v1/answers/:id/flag",    "Flag an answer for scholar review"),
    ("GET",  "/api/v1/quran/:surah/:ayah",  "Fetch a specific Quran verse with translations"),
    ("GET",  "/api/v1/hadith/:collection/:id", "Fetch a specific hadith by collection and number"),
    ("GET",  "/api/v1/search",              "Full-text keyword search across Quran and Hadith"),
    ("GET",  "/api/v1/daily",               "Get today's Hadith of the Day and Verse of the Day"),
    ("POST", "/api/v1/users/register",      "Create a user account"),
    ("POST", "/api/v1/users/login",         "Authenticate and receive JWT"),
    ("GET",  "/api/v1/users/me/saved",      "Retrieve saved Q&A items for authenticated user"),
    ("POST", "/api/v1/users/me/saved",      "Save a Q&A item"),
    ("GET",  "/api/v1/scholar/flags",       "Scholar: list flagged answers for review (scholar role required)"),
    ("PUT",  "/api/v1/scholar/flags/:id",   "Scholar: approve, reject, or annotate a flagged answer"),
]

tbl7 = doc.add_table(rows=1, cols=3)
tbl7.style = 'Table Grid'
tbl7.columns[0].width = Inches(0.7)
tbl7.columns[1].width = Inches(3.0)
tbl7.columns[2].width = Inches(3.55)
add_table_row(tbl7, ["Method", "Endpoint", "Description"], header=True)
for r in endpoints:
    add_table_row(tbl7, list(r))

# ═══════════════════════════════════════════════════════════════════════════════
# 11. DEVELOPMENT PHASES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("11. Development Phases & Milestones")

phases = [
    ("Phase 1 — Foundation (Weeks 1–4)",
     [
         "Set up monorepo (Next.js web + React Native Expo mobile).",
         "Provision PostgreSQL + pgvector database; design and run migrations for all core tables.",
         "Ingest Quran data (all 6,236 ayat, Arabic + EN/UR translations) and generate embeddings.",
         "Ingest Sahih Bukhari + Sahih Muslim as MVP hadith sources; grade-filter to Sahih only.",
         "Build and test the RAG pipeline: embedding search → context retrieval → LLM call → citation verification.",
         "Basic REST API: /ask, /quran, /hadith, /search endpoints.",
         "Internal demo with 50 test questions evaluated by a scholar reviewer.",
     ]
    ),
    ("Phase 2 — Core App (Weeks 5–9)",
     [
         "Build mobile and web UI: home screen, question input, answer display with citations, source cards.",
         "RTL layout support for Arabic and Urdu.",
         "User authentication (email + Google/Apple sign-in).",
         "Save / history / share functionality.",
         "Ingest remaining 6 Hadith collections (Sunan Abu Dawud, Tirmidhi, Nasa'i, Ibn Majah, Muwatta, Musnad Ahmad) with daif filtering.",
         "Scholar review dashboard (web only in this phase).",
         "Push notifications for Hadith of the Day.",
     ]
    ),
    ("Phase 3 — Languages & Quality (Weeks 10–13)",
     [
         "Add French, Turkish, Indonesian, Bangla as Phase 2 languages.",
         "Language auto-detection and per-language prompt tuning.",
         "Improve confidence scoring and escalation logic.",
         "Run red-team evaluation: 500 adversarial questions to test hallucination safeguards.",
         "Performance optimisation: response time target < 3 seconds for 95th percentile.",
         "Beta programme: invite 200 community members (including scholars) for feedback.",
     ]
    ),
    ("Phase 4 — Launch & Growth (Weeks 14–16)",
     [
         "App Store (iOS) and Google Play Store submission and approval.",
         "PWA launch on web with SEO optimisation.",
         "Public launch marketing to Muslim community forums, mosques, Islamic centres.",
         "Monitoring dashboards live; on-call rotation for backend team.",
         "Collect and triage user feedback; plan v1.1 feature backlog.",
     ]
    ),
]

for phase_name, tasks in phases:
    add_heading(phase_name, level=2)
    for t in tasks:
        add_bullet(t)

# ═══════════════════════════════════════════════════════════════════════════════
# 12. SECURITY & PRIVACY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("12. Security & Privacy")
add_bullet("All API traffic over HTTPS/TLS 1.3. No plain-text transmission.")
add_bullet("User passwords hashed with bcrypt (min 12 rounds). JWT tokens short-lived (15 min access + 7-day refresh).")
add_bullet("Questions submitted by guest users are stored anonymously (no IP logged after 24 hours).")
add_bullet("LLM API calls made server-side only — API keys never exposed to clients.")
add_bullet("Input sanitisation and rate limiting on all endpoints (100 req/min per IP for /ask).")
add_bullet("GDPR & data-minimisation compliance: users can delete their account and all associated data.")
add_bullet("Scholar accounts require email verification + manual approval by admin.")
add_bullet("Content moderation layer to prevent prompt injection attacks attempting to override system prompt.")

# ═══════════════════════════════════════════════════════════════════════════════
# 13. SCHOLAR & COMMUNITY TRUST STRATEGY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("13. Scholar & Community Trust Strategy")
add_body(
    "Trust is the single most important factor for a faith-based application. The following "
    "measures should be implemented before public launch:"
)
add_bullet("Recruit a panel of at least 3 qualified Islamic scholars (from different madhabs/schools of thought) to review the system before launch and serve as ongoing advisors.")
add_bullet("Publish a clear 'About Our Sources' page listing every Hadith collection used, the grading methodology applied, and the names of the scholar advisory panel.")
add_bullet("Every answer page must display a disclaimer: 'This response is generated by AI and is for educational purposes only. It is not a fatwa. Please consult a qualified scholar for religious rulings.'")
add_bullet("Implement a 'Ask a Scholar' escalation button that emails the question to the advisory panel for complex matters.")
add_bullet("Publish the grading filter criteria publicly so the community can audit which hadith are included.")
add_bullet("Regular (quarterly) scholar audits of a random sample of AI answers.")

# ═══════════════════════════════════════════════════════════════════════════════
# 14. MONETISATION (OPTIONAL)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("14. Monetisation Options (Optional)")
add_body("The app can be free for all core features. Optional sustainable revenue streams include:")
add_bullet("Voluntary donation / Sadaqah button — position as community service, not a commercial product.")
add_bullet("NoorAI Premium subscription: offline full corpus download, ad-free, extended daily reminders, family accounts.")
add_bullet("Institutional licensing: mosques, Islamic schools, and universities pay for a white-labelled version.")
add_bullet("No advertising — maintaining the app as a clean, focused spiritual tool is essential for community trust.")

# ═══════════════════════════════════════════════════════════════════════════════
# 15. ESTIMATED EFFORT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("15. Estimated Build Effort")

tbl8 = doc.add_table(rows=1, cols=3)
tbl8.style = 'Table Grid'
tbl8.columns[0].width = Inches(2.5)
tbl8.columns[1].width = Inches(1.5)
tbl8.columns[2].width = Inches(3.25)
add_table_row(tbl8, ["Area", "Est. Weeks", "Key Deliverable"], header=True)
effort_rows = [
    ("Knowledge Base Ingestion & RAG Pipeline", "3–4",  "Quran + all Hadith collections embedded and searchable"),
    ("Backend API + Auth",                       "3–4",  "All REST endpoints live, tested, documented"),
    ("Web App (Next.js PWA)",                    "3–4",  "Full-featured web app deployed on Vercel/CloudFlare"),
    ("Mobile App (React Native)",                "4–5",  "iOS + Android apps passing App Store review"),
    ("Multilingual & RTL",                       "2–3",  "EN, AR, UR live; RTL layout correct"),
    ("Scholar Review Dashboard",                 "1–2",  "Flagging, annotation, approval workflow live"),
    ("QA, Security Audit, Beta Test",            "2–3",  "Zero critical issues; scholar sign-off"),
    ("TOTAL",                                    "16–18 weeks", "Full production launch"),
]
for r in effort_rows:
    add_table_row(tbl8, list(r))

doc.add_paragraph()
add_note("This estimate assumes a team of 2–3 full-stack developers. A solo developer should plan for 24–30 weeks.")

# ═══════════════════════════════════════════════════════════════════════════════
# 16. INSTRUCTIONS FOR CLAUDE (HOW TO BUILD)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading("16. Instructions for Claude — How to Build This App")
add_body(
    "When handing this document to Claude to build the application, provide the following "
    "context and instructions in your prompt:"
)

add_heading("16.1  Recommended Prompt to Give Claude", level=2)
p_block = doc.add_paragraph()
p_block.paragraph_format.left_indent  = Inches(0.4)
p_block.paragraph_format.space_before = Pt(4)
p_block.paragraph_format.space_after  = Pt(4)
# light grey background via paragraph shading
pPr2 = p_block._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'),   'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'),  'F2F2F2')
pPr2.append(shd2)
prompt_text = (
    '"Please build the NoorAI Islamic Q&A application as described in the attached build plan. '
    'Start with Phase 1: (1) Set up a monorepo with Next.js for web and React Native Expo for mobile. '
    '(2) Create the PostgreSQL schema for quran_verses and hadith tables. '
    '(3) Write a data ingestion script to load the Quran (use quran.com API) and Sahih Bukhari/Muslim (use sunnah.com API or hadith-json on GitHub). '
    '(4) Generate embeddings using OpenAI text-embedding-3-small and store them in pgvector. '
    '(5) Build the /api/v1/ask RAG endpoint using LangChain and the Claude API, with the strict grounding system prompt from Section 8 of the plan. '
    '(6) Add the citation verification post-processing step. '
    'Please work phase by phase and confirm with me after each phase before proceeding to the next."'
)
r_prompt = p_block.add_run(prompt_text)
set_run_font(r_prompt, size=10, italic=True, color=RGBColor(0x33, 0x33, 0x33))

add_heading("16.2  Key Files & Folders Claude Should Create", level=2)
add_bullet("apps/web/ — Next.js PWA")
add_bullet("apps/mobile/ — React Native Expo app")
add_bullet("packages/shared/ — shared TypeScript types, API client, i18n strings")
add_bullet("services/api/ — Fastify or FastAPI backend")
add_bullet("services/ingestion/ — Quran + Hadith data ingestion scripts")
add_bullet("services/rag/ — RAG pipeline (embedding search, LLM call, citation verifier)")
add_bullet("database/migrations/ — PostgreSQL schema migrations")
add_bullet("docs/ — API documentation, scholar guidelines, source audit log")

add_heading("16.3  Environment Variables Claude Will Need", level=2)
add_bullet("ANTHROPIC_API_KEY — Claude LLM API key")
add_bullet("OPENAI_API_KEY — for text embeddings")
add_bullet("DATABASE_URL — PostgreSQL connection string")
add_bullet("REDIS_URL — Redis connection string")
add_bullet("SUPABASE_URL + SUPABASE_ANON_KEY — if using Supabase for auth/storage")
add_bullet("SUNNAH_COM_API_KEY — for Hadith data ingestion")
add_bullet("PINECONE_API_KEY + PINECONE_ENV — if using Pinecone instead of pgvector")

# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_close = closing.add_run("— May this app be a source of benefit (nafa) for the Ummah — Ameen —")
set_run_font(r_close, size=11, italic=True, color=GOLD)

# ── Save ───────────────────────────────────────────────────────────────────────
out = "/sessions/intelligent-blissful-noether/mnt/outputs/NoorAI_Build_Plan.docx"
doc.save(out)
print("Saved:", out)
